"""Extracts text from a stored document file: .txt, .md, .pdf, .docx, and .xlsx only.

No chunking, embedding, or Qdrant upsert here — this is purely the "load the file and get raw
text out of it" step. Routing is by file extension, but before parsing, each file's basic
structure/content is validated against what its extension claims (PDF header, DOCX/XLSX OOXML
zip structure, UTF-8 readability for plain text) — a mismatched or corrupt file fails clearly
instead of being handed to the wrong parser. PDF text is extracted page by page via pypdf,
preserving page numbers; XLSX is extracted sheet by sheet via openpyxl, preserving sheet names;
DOCX is extracted as plain paragraph text via python-docx; plain text/Markdown files are treated
as a single unnumbered page.
"""

import asyncio
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
from openpyxl import load_workbook
from pypdf import PdfReader

from app.models.document import Document

_PLAIN_TEXT_SUFFIXES = {".txt", ".md"}
_SUPPORTED_SUFFIXES = _PLAIN_TEXT_SUFFIXES | {".pdf", ".docx", ".xlsx"}
_PDF_HEADER = b"%PDF"
_DOCX_REQUIRED_ENTRY = "word/document.xml"
_XLSX_REQUIRED_ENTRY = "xl/workbook.xml"


class DocumentTextExtractionError(Exception):
    """Raised when a document's stored file is missing, unsupported, or has no extractable text."""


@dataclass
class ExtractedPage:
    """One page's extracted text.

    `page_number` is set for PDFs, `sheet_name` for XLSX sheets — both None for plain
    text/Markdown/DOCX, which have no natural pagination.
    """

    text: str
    page_number: int | None = None
    sheet_name: str | None = None


@dataclass
class ExtractedDocument:
    """All extracted text for one Document, as an ordered list of pages."""

    document_id: str
    pages: list[ExtractedPage]

    @property
    def full_text(self) -> str:
        """Return all pages' text concatenated in order, separated by newlines."""
        return "\n".join(page.text for page in self.pages)


class DocumentTextExtractor:
    """Extracts text from a Document's stored file into an ExtractedDocument."""

    async def extract(self, document: Document) -> ExtractedDocument:
        """Load the document's stored file and extract its text (blocking work off the loop)."""
        return await asyncio.to_thread(self._extract_sync, document)

    def _extract_sync(self, document: Document) -> ExtractedDocument:
        path = Path(document.stored_path)
        if not path.exists():
            raise DocumentTextExtractionError(f"Stored file not found: {document.stored_path}")

        suffix = path.suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            raise DocumentTextExtractionError(f"Unsupported file type: {suffix or '(no extension)'}")

        self._validate_file_type(path, suffix)

        if suffix in _PLAIN_TEXT_SUFFIXES:
            pages = [self._extract_plain_text(path)]
        elif suffix == ".pdf":
            pages = self._extract_pdf(path)
        elif suffix == ".docx":
            pages = [self._extract_docx(path)]
        else:
            pages = self._extract_xlsx(path)

        if not any(page.text.strip() for page in pages):
            raise DocumentTextExtractionError("No extractable text found in document.")

        return ExtractedDocument(document_id=document.id, pages=pages)

    @staticmethod
    def _validate_file_type(path: Path, suffix: str) -> None:
        """Check the file's actual content matches what its extension claims, before parsing it."""
        if suffix in _PLAIN_TEXT_SUFFIXES:
            try:
                path.read_text(encoding="utf-8")
            except UnicodeDecodeError as exc:
                raise DocumentTextExtractionError(f"File is not valid UTF-8 text: {path.name}") from exc
            return

        if suffix == ".pdf":
            with path.open("rb") as f:
                header = f.read(len(_PDF_HEADER))
            if header != _PDF_HEADER:
                raise DocumentTextExtractionError(f"File does not look like a valid PDF: {path.name}")
            return

        if suffix == ".docx":
            DocumentTextExtractor._validate_office_zip(path, _DOCX_REQUIRED_ENTRY, "DOCX")
            return

        DocumentTextExtractor._validate_office_zip(path, _XLSX_REQUIRED_ENTRY, "XLSX")

    @staticmethod
    def _validate_office_zip(path: Path, required_entry: str, kind: str) -> None:
        """Check the file is a valid ZIP archive containing the expected OOXML structure."""
        if not zipfile.is_zipfile(path):
            raise DocumentTextExtractionError(
                f"File does not look like a valid {kind} (not a zip archive): {path.name}"
            )
        with zipfile.ZipFile(path) as archive:
            if required_entry not in archive.namelist():
                raise DocumentTextExtractionError(
                    f"File does not look like a valid {kind} (missing {required_entry}): {path.name}"
                )

    @staticmethod
    def _extract_plain_text(path: Path) -> ExtractedPage:
        """Read a .txt/.md file as UTF-8 text (Hebrew and other Unicode content supported)."""
        text = path.read_text(encoding="utf-8")
        return ExtractedPage(text=text, page_number=None)

    @staticmethod
    def _extract_pdf(path: Path) -> list[ExtractedPage]:
        """Extract text page by page from a PDF, preserving 1-indexed page numbers."""
        reader = PdfReader(str(path))
        return [
            ExtractedPage(text=page.extract_text() or "", page_number=index + 1)
            for index, page in enumerate(reader.pages)
        ]

    @staticmethod
    def _extract_docx(path: Path) -> ExtractedPage:
        """Extract plain paragraph text from a DOCX (no tables, headers/footers, or pagination)."""
        document = docx.Document(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ExtractedPage(text=text)

    @staticmethod
    def _extract_xlsx(path: Path) -> list[ExtractedPage]:
        """Extract text sheet by sheet from an XLSX, preserving each sheet's name."""
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        pages = []
        for sheet in workbook.worksheets:
            rows_text = [
                "\t".join(str(cell) for cell in row if cell is not None)
                for row in sheet.iter_rows(values_only=True)
                if any(cell is not None for cell in row)
            ]
            pages.append(ExtractedPage(text="\n".join(rows_text), sheet_name=sheet.title))
        return pages
