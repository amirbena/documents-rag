"""Extracts text from a stored document's bytes: .txt, .md, .pdf, .docx, and .xlsx only.

No chunking, embedding, or Qdrant upsert here — this is purely the "load the bytes and get raw
text out of it" step. Content is read via the injected `FileStorage` (never a direct filesystem
path) and parsed entirely in memory — pypdf/python-docx/openpyxl all accept an in-memory
`io.BytesIO` stream directly, so no temporary file materialization is needed. Routing is by file
extension (from the document's original filename), but before parsing, each file's basic
structure/content is validated against what its extension claims (PDF header, DOCX/XLSX OOXML
zip structure, UTF-8 readability for plain text) — a mismatched or corrupt file fails clearly
instead of being handed to the wrong parser. PDF text is extracted page by page via pypdf,
preserving page numbers; XLSX is extracted sheet by sheet via openpyxl, preserving sheet names;
DOCX is extracted as plain paragraph text via python-docx; plain text/Markdown files are treated
as a single unnumbered page.
"""

import asyncio
import io
import zipfile
from dataclasses import dataclass
from pathlib import Path

import docx
from openpyxl import load_workbook
from pypdf import PdfReader

from app.models.document import Document
from app.storage.contract import FileStorage
from app.storage.errors import StorageObjectNotFoundError
from app.storage.keys import resolve_document_storage_key

_PLAIN_TEXT_SUFFIXES = {".txt", ".md"}
_SUPPORTED_SUFFIXES = _PLAIN_TEXT_SUFFIXES | {".pdf", ".docx", ".xlsx"}
_PDF_HEADER = b"%PDF"
_DOCX_REQUIRED_ENTRY = "word/document.xml"
_XLSX_REQUIRED_ENTRY = "xl/workbook.xml"


# Category (Phase 2.10, see app/core/errors.py): ValidationError — the stored content doesn't
# match what ingestion needs to proceed.


class DocumentTextExtractionError(Exception):
    """Raised when a document's stored content is missing, unsupported, or has no extractable text."""


@dataclass
class ExtractedPage:
    """One page's extracted text.

    `page_number` is set for PDFs, `sheet_name` for XLSX sheets — both None for plain
    text/Markdown/DOCX, which have no natural pagination.
    """

    text: str
    page_number: int | None = None
    sheet_name: str | None = None


@dataclass
class ExtractedDocument:
    """All extracted text for one Document, as an ordered list of pages."""

    document_id: str
    pages: list[ExtractedPage]

    @property
    def full_text(self) -> str:
        """Return all pages' text concatenated in order, separated by newlines."""
        return "\n".join(page.text for page in self.pages)


class DocumentTextExtractor:
    """Extracts text from a Document's stored content (read via FileStorage) into an ExtractedDocument."""

    def __init__(self, storage: FileStorage) -> None:
        self._storage = storage

    async def extract(self, document: Document) -> ExtractedDocument:
        """Read the document's stored content via FileStorage and extract its text."""
        key = resolve_document_storage_key(document)
        try:
            content = await self._storage.read(key)
        except StorageObjectNotFoundError as exc:
            raise DocumentTextExtractionError(f"Stored file not found: {key}") from exc

        return await asyncio.to_thread(self._extract_sync, document, content)

    def _extract_sync(self, document: Document, content: bytes) -> ExtractedDocument:
        suffix = Path(document.original_filename).suffix.lower()
        if suffix not in _SUPPORTED_SUFFIXES:
            raise DocumentTextExtractionError(f"Unsupported file type: {suffix or '(no extension)'}")

        self._validate_file_type(content, suffix, document.original_filename)

        if suffix in _PLAIN_TEXT_SUFFIXES:
            pages = [self._extract_plain_text(content)]
        elif suffix == ".pdf":
            pages = self._extract_pdf(content)
        elif suffix == ".docx":
            pages = [self._extract_docx(content)]
        else:
            pages = self._extract_xlsx(content)

        if not any(page.text.strip() for page in pages):
            raise DocumentTextExtractionError("No extractable text found in document.")

        return ExtractedDocument(document_id=document.id, pages=pages)

    @staticmethod
    def _validate_file_type(content: bytes, suffix: str, filename: str) -> None:
        """Check the file's actual content matches what its extension claims, before parsing it."""
        if suffix in _PLAIN_TEXT_SUFFIXES:
            try:
                content.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise DocumentTextExtractionError(f"File is not valid UTF-8 text: {filename}") from exc
            return

        if suffix == ".pdf":
            if content[: len(_PDF_HEADER)] != _PDF_HEADER:
                raise DocumentTextExtractionError(f"File does not look like a valid PDF: {filename}")
            return

        if suffix == ".docx":
            DocumentTextExtractor._validate_office_zip(content, _DOCX_REQUIRED_ENTRY, "DOCX", filename)
            return

        DocumentTextExtractor._validate_office_zip(content, _XLSX_REQUIRED_ENTRY, "XLSX", filename)

    @staticmethod
    def _validate_office_zip(content: bytes, required_entry: str, kind: str, filename: str) -> None:
        """Check the file is a valid ZIP archive containing the expected OOXML structure."""
        buffer = io.BytesIO(content)
        if not zipfile.is_zipfile(buffer):
            raise DocumentTextExtractionError(
                f"File does not look like a valid {kind} (not a zip archive): {filename}"
            )
        buffer.seek(0)
        with zipfile.ZipFile(buffer) as archive:
            if required_entry not in archive.namelist():
                raise DocumentTextExtractionError(
                    f"File does not look like a valid {kind} (missing {required_entry}): {filename}"
                )

    @staticmethod
    def _extract_plain_text(content: bytes) -> ExtractedPage:
        """Decode a .txt/.md file's bytes as UTF-8 text (Hebrew and other Unicode content supported)."""
        return ExtractedPage(text=content.decode("utf-8"), page_number=None)

    @staticmethod
    def _extract_pdf(content: bytes) -> list[ExtractedPage]:
        """Extract text page by page from a PDF, preserving 1-indexed page numbers."""
        reader = PdfReader(io.BytesIO(content))
        return [
            ExtractedPage(text=page.extract_text() or "", page_number=index + 1)
            for index, page in enumerate(reader.pages)
        ]

    @staticmethod
    def _extract_docx(content: bytes) -> ExtractedPage:
        """Extract plain paragraph text from a DOCX (no tables, headers/footers, or pagination)."""
        document = docx.Document(io.BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ExtractedPage(text=text)

    @staticmethod
    def _extract_xlsx(content: bytes) -> list[ExtractedPage]:
        """Extract text sheet by sheet from an XLSX, preserving each sheet's name."""
        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        pages = []
        for sheet in workbook.worksheets:
            rows_text = [
                "\t".join(str(cell) for cell in row if cell is not None)
                for row in sheet.iter_rows(values_only=True)
                if any(cell is not None for cell in row)
            ]
            pages.append(ExtractedPage(text="\n".join(rows_text), sheet_name=sheet.title))
        return pages
