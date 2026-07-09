"""Tests for DocumentTextExtractor: .txt/.md/.pdf/.docx/.xlsx extraction, errors, Unicode text."""

import io
import uuid
import zipfile
from pathlib import Path

import docx
import pytest
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DictionaryObject, NameObject
from pypdf.generic import StreamObject as _StreamObject

from app.models.document import Document
from app.services.document_text_extractor import DocumentTextExtractionError, DocumentTextExtractor


def _make_document(path: Path, content_type: str = "text/plain") -> Document:
    return Document(
        id=str(uuid.uuid4()),
        original_filename=path.name,
        stored_filename=path.name,
        content_type=content_type,
        file_size=path.stat().st_size if path.exists() else 0,
        stored_path=str(path),
    )


def _build_pdf_bytes(page_texts: list[str]) -> bytes:
    """Build a minimal multi-page PDF with real extractable text, using pypdf only.

    pypdf has no public "set page content" API, so this uses `_add_object` (private) to attach
    a raw content stream to a blank page — a known idiom for building pypdf test fixtures
    without a separate PDF-writing dependency. Test-fixture generation only, not production code.
    """
    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=200, height=200)

        stream = _StreamObject()
        stream.set_data(f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode("latin-1"))
        stream_ref = writer._add_object(stream)
        page[NameObject("/Contents")] = stream_ref

        font = DictionaryObject()
        font[NameObject("/Type")] = NameObject("/Font")
        font[NameObject("/Subtype")] = NameObject("/Type1")
        font[NameObject("/BaseFont")] = NameObject("/Helvetica")
        font_ref = writer._add_object(font)

        fonts = DictionaryObject()
        fonts[NameObject("/F1")] = font_ref
        resources = DictionaryObject()
        resources[NameObject("/Font")] = fonts
        page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


async def test_txt_extraction(tmp_path: Path) -> None:
    """A .txt file should extract as a single unnumbered page."""
    path = tmp_path / "notes.txt"
    path.write_text("hello world", encoding="utf-8")
    document = _make_document(path)

    result = await DocumentTextExtractor().extract(document)

    assert result.document_id == document.id
    assert len(result.pages) == 1
    assert result.pages[0].text == "hello world"
    assert result.pages[0].page_number is None


async def test_markdown_extraction(tmp_path: Path) -> None:
    """A .md file should extract as a single unnumbered page, raw text (no markdown parsing)."""
    path = tmp_path / "readme.md"
    path.write_text("# Title\n\nSome **bold** text.", encoding="utf-8")
    document = _make_document(path, content_type="text/markdown")

    result = await DocumentTextExtractor().extract(document)

    assert len(result.pages) == 1
    assert result.pages[0].text == "# Title\n\nSome **bold** text."
    assert result.pages[0].page_number is None


async def test_pdf_extraction_with_page_numbers(tmp_path: Path) -> None:
    """A PDF should extract text page by page, with 1-indexed page numbers preserved."""
    path = tmp_path / "handbook.pdf"
    path.write_bytes(_build_pdf_bytes(["Page one text", "Page two text"]))
    document = _make_document(path, content_type="application/pdf")

    result = await DocumentTextExtractor().extract(document)

    assert len(result.pages) == 2
    assert result.pages[0].page_number == 1
    assert "Page one text" in result.pages[0].text
    assert result.pages[1].page_number == 2
    assert "Page two text" in result.pages[1].text
    assert "Page one text" in result.full_text
    assert "Page two text" in result.full_text


async def test_unsupported_file_type_fails(tmp_path: Path) -> None:
    """An unsupported extension should raise DocumentTextExtractionError."""
    path = tmp_path / "archive.zip"
    path.write_bytes(b"not a real zip, just bytes")
    document = _make_document(path, content_type="application/zip")

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_missing_file_fails(tmp_path: Path) -> None:
    """A stored_path that doesn't exist on disk should raise DocumentTextExtractionError."""
    path = tmp_path / "does_not_exist.txt"
    document = _make_document(path)

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_empty_extracted_text_fails(tmp_path: Path) -> None:
    """A file with no meaningful text content should raise DocumentTextExtractionError."""
    path = tmp_path / "blank.txt"
    path.write_text("   \n\n  ", encoding="utf-8")
    document = _make_document(path)

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_hebrew_text_extraction(tmp_path: Path) -> None:
    """A UTF-8 .txt file with Hebrew content should extract correctly, preserved exactly."""
    path = tmp_path / "חוזה.txt"
    hebrew_text = "שלום עולם, זהו מסמך לדוגמה."
    path.write_text(hebrew_text, encoding="utf-8")
    document = _make_document(path)

    result = await DocumentTextExtractor().extract(document)

    assert result.pages[0].text == hebrew_text


async def test_docx_extraction(tmp_path: Path) -> None:
    """A .docx file should extract as a single unnumbered page of plain paragraph text."""
    path = tmp_path / "letter.docx"
    document_file = docx.Document()
    document_file.add_paragraph("First paragraph.")
    document_file.add_paragraph("Second paragraph.")
    document_file.save(str(path))
    document = _make_document(
        path, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    result = await DocumentTextExtractor().extract(document)

    assert len(result.pages) == 1
    assert result.pages[0].page_number is None
    assert result.pages[0].sheet_name is None
    assert "First paragraph." in result.pages[0].text
    assert "Second paragraph." in result.pages[0].text


async def test_docx_hebrew_extraction(tmp_path: Path) -> None:
    """A .docx file with Hebrew content should extract correctly, preserved exactly."""
    path = tmp_path / "מכתב.docx"
    hebrew_text = "שלום עולם, זהו מסמך וורד לדוגמה."
    document_file = docx.Document()
    document_file.add_paragraph(hebrew_text)
    document_file.save(str(path))
    document = _make_document(path)

    result = await DocumentTextExtractor().extract(document)

    assert hebrew_text in result.pages[0].text


async def test_xlsx_extraction(tmp_path: Path) -> None:
    """An .xlsx file should extract as one page per sheet, preserving the sheet name."""
    path = tmp_path / "budget.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.append(["Item", "Amount"])
    sheet.append(["Rent", 1000])
    workbook.save(str(path))
    document = _make_document(
        path, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    result = await DocumentTextExtractor().extract(document)

    assert len(result.pages) == 1
    assert result.pages[0].sheet_name == "Summary"
    assert result.pages[0].page_number is None
    assert "Item" in result.pages[0].text
    assert "Rent" in result.pages[0].text


async def test_xlsx_multiple_sheets(tmp_path: Path) -> None:
    """An .xlsx file with multiple sheets should extract one page per sheet, in order."""
    path = tmp_path / "report.xlsx"
    workbook = Workbook()
    first_sheet = workbook.active
    first_sheet.title = "Q1"
    first_sheet.append(["Revenue", 100])
    second_sheet = workbook.create_sheet("Q2")
    second_sheet.append(["Revenue", 200])
    workbook.save(str(path))
    document = _make_document(path)

    result = await DocumentTextExtractor().extract(document)

    assert len(result.pages) == 2
    assert result.pages[0].sheet_name == "Q1"
    assert "100" in result.pages[0].text
    assert result.pages[1].sheet_name == "Q2"
    assert "200" in result.pages[1].text


async def test_xlsx_hebrew_extraction(tmp_path: Path) -> None:
    """An .xlsx file with Hebrew cell content should extract correctly, preserved exactly."""
    path = tmp_path / "תקציב.xlsx"
    hebrew_value = "שכירות"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append([hebrew_value, 1000])
    workbook.save(str(path))
    document = _make_document(path)

    result = await DocumentTextExtractor().extract(document)

    assert hebrew_value in result.pages[0].text


async def test_fake_pdf_with_non_pdf_bytes_fails(tmp_path: Path) -> None:
    """A .pdf file whose bytes don't start with the %PDF header should fail validation."""
    path = tmp_path / "fake.pdf"
    path.write_bytes(b"this is not actually a pdf, just plain bytes")
    document = _make_document(path, content_type="application/pdf")

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_fake_docx_with_non_docx_bytes_fails(tmp_path: Path) -> None:
    """A .docx file whose bytes aren't a valid OOXML zip should fail validation."""
    path = tmp_path / "fake.docx"
    path.write_bytes(b"this is not a real docx, just plain bytes")
    document = _make_document(
        path, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_fake_xlsx_with_non_xlsx_bytes_fails(tmp_path: Path) -> None:
    """An .xlsx file whose bytes aren't a valid OOXML zip should fail validation."""
    path = tmp_path / "fake.xlsx"
    path.write_bytes(b"this is not a real xlsx, just plain bytes")
    document = _make_document(
        path, content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)


async def test_zip_without_docx_structure_fails_as_docx(tmp_path: Path) -> None:
    """A valid zip file that isn't a real DOCX (missing word/document.xml) should still fail."""
    path = tmp_path / "fake.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("not_a_docx_entry.txt", "hello")
    document = _make_document(path)

    with pytest.raises(DocumentTextExtractionError):
        await DocumentTextExtractor().extract(document)
